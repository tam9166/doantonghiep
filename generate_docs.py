import sys
import os
import subprocess

# 1. Tự động kiểm tra và cài đặt python-docx nếu chưa có
try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn
except ImportError:
    print("Không tìm thấy python-docx. Đang tiến hành cài đặt tự động...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn

# --- CÁC HÀM TRỢ GIÚP ĐỂ TẠO STYLE CHO FILE WORD ĐẸP VÀ CHUYÊN NGHIỆP ---

def set_cell_background(cell, color_hex):
    """Đặt màu nền cho cell trong bảng"""
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Đặt khoảng cách đệm (padding) cho cell"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_styled(doc, text, level):
    """Thêm tiêu đề được thiết kế đẹp (Font Times New Roman, Màu Xanh Đậm, Đậm)"""
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    
    # Định dạng font cho tiêu đề
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.bold = True
        if level == 1:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(30, 58, 138) # Dark Blue #1E3A8A
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(51, 65, 85) # Slate #334155
        else:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(100, 116, 139)
    return p

def add_paragraph_styled(doc, text="", bold_prefix="", italic=False):
    """Thêm đoạn văn chuẩn Times New Roman 12pt, dãn dòng 1.25"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(6)
    
    if bold_prefix:
        r_prefix = p.add_run(bold_prefix)
        r_prefix.font.name = 'Times New Roman'
        r_prefix.font.size = Pt(12)
        r_prefix.font.bold = True
        
    if text:
        r_text = p.add_run(text)
        r_text.font.name = 'Times New Roman'
        r_text.font.size = Pt(12)
        r_text.font.italic = italic
        
    return p

def add_bullet_styled(doc, text, bold_prefix=""):
    """Thêm gạch đầu dòng chuẩn đẹp"""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(4)
    
    if bold_prefix:
        r_prefix = p.add_run(bold_prefix)
        r_prefix.font.name = 'Times New Roman'
        r_prefix.font.size = Pt(12)
        r_prefix.font.bold = True
        
    r_text = p.add_run(text)
    r_text.font.name = 'Times New Roman'
    r_text.font.size = Pt(12)
    return p

def setup_document_base(title_text):
    """Khởi tạo tài liệu và thiết lập tiêu đề chính trang bìa"""
    doc = docx.Document()
    
    # Set Margins (Chuẩn Việt Nam: Lề trái 3cm, lề khác 2cm)
    for section in doc.sections:
        section.top_margin = Inches(0.79) # 2cm
        section.bottom_margin = Inches(0.79) # 2cm
        section.left_margin = Inches(1.18) # 3cm
        section.right_margin = Inches(0.79) # 2cm
        
    # Thêm tiêu đề lớn
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(18)
    p_title.paragraph_format.space_after = Pt(18)
    
    run_title = p_title.add_run(title_text)
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(30, 58, 138)
    
    return doc

def save_document(doc, filename):
    """Lưu tài liệu xuống đĩa"""
    filepath = os.path.join("d:\\DoAnTotNghiep", filename)
    doc.save(filepath)
    print(f"Đã lưu file thành công tại: {filepath}")

# ==============================================================================
# FILE 1: MÔ TẢ HỆ THỐNG (MotaHeThong.docx)
# ==============================================================================
def generate_file_mota():
    doc = setup_document_base("TÀI LIỆU MÔ TẢ HỆ THỐNG VÀ CÁC MÀN HÌNH\nHỆ THỐNG QUẢN LÝ NHÀ HÀNG MỘC VỊ RESTAURANT")
    
    add_heading_styled(doc, "1. Giới thiệu tổng quan hệ thống Mộc Vị Restaurant", level=1)
    add_paragraph_styled(doc, "Hệ thống Quản lý Nhà hàng Mộc Vị là một giải pháp chuyển đổi số toàn diện cho mô hình nhà hàng ăn uống, bao gồm sự kết hợp chặt chẽ giữa vận hành tiền sảnh (gọi món tại bàn, đặt bàn, thanh toán nhanh) và hậu cần nhà bếp (tiếp nhận món ăn theo thời gian thực, quản lý kho nguyên liệu và định lượng công thức).")
    add_paragraph_styled(doc, "Điểm nổi bật của hệ thống là việc tích hợp công nghệ AI (Gemini API) hỗ trợ nhân viên phục vụ tư vấn bán chéo món ăn (cross-selling) và thanh toán quét mã VietQR tự động tạo động cho từng bàn, giúp tối ưu quy trình phục vụ và gia tăng đáng kể doanh thu cho nhà hàng.")
    
    add_heading_styled(doc, "2. Các vai trò người dùng (Roles & Actors)", level=1)
    add_paragraph_styled(doc, "Hệ thống phân chia quyền truy cập chặt chẽ thành 4 vai trò chính thông qua cấu hình Spring Security ở Backend:")
    
    add_bullet_styled(doc, " Xem thực đơn số, đặt bàn hẹn giờ trước khi đến, quét mã QR tại bàn để gọi món tự động, theo dõi trạng thái đơn hàng và gửi phản hồi đánh giá chất lượng món ăn.", bold_prefix="Khách hàng (Guest / Customer):")
    add_bullet_styled(doc, " Quản lý sơ đồ bàn trực quan (bàn trống, bàn có khách, bàn đang dọn), hỗ trợ khách gọi thêm món, thực hiện chuyển bàn ăn tự động, in hóa đơn tạm tính kèm mã QR thanh toán động và nhận ý kiến gợi ý món từ trợ lý AI để upsell.", bold_prefix="Nhân viên Phục vụ (Waiter):")
    add_bullet_styled(doc, " Nhận danh sách món ăn cần chế biến tức thời theo thời gian thực qua kết nối WebSockets, cập nhật trạng thái chế biến (đang nấu, xong món) và thông báo đóng/mở bán món ăn tùy theo tình trạng kho nguyên liệu.", bold_prefix="Nhân viên Bếp (Kitchen):")
    add_bullet_styled(doc, " Nắm toàn quyền quản trị hệ thống, bao gồm: quản lý tài khoản nhân viên và phân quyền, quản lý danh mục và thông tin món ăn, quản lý kho nguyên liệu & công thức chế biến (Recipes), xem bảng thống kê doanh thu và sản phẩm bán chạy qua biểu đồ trực quan.", bold_prefix="Quản lý / Admin (Admin / Manager):")

    add_heading_styled(doc, "3. Các chức năng cốt lõi của dự án", level=1)
    
    add_heading_styled(doc, "3.1. Gọi món trực tuyến tại bàn & Đặt bàn hẹn giờ", level=2)
    add_paragraph_styled(doc, "Khách hàng có thể đặt bàn trước thông qua giao diện Web công cộng, tùy chọn số lượng người, thời gian hẹn và vị trí bàn mong muốn. Khi ngồi tại bàn ăn thực tế, khách quét mã QR để truy cập thực đơn số của riêng bàn đó, tiến hành chọn món gửi thẳng xuống bếp mà không cần gọi phục vụ.")
    
    add_heading_styled(doc, "3.2. Đồng bộ bếp - phục vụ thời gian thực (Real-time WebSockets)", level=2)
    add_paragraph_styled(doc, "Nhờ tích hợp giao thức STOMP trên nền tảng WebSocket, mọi hành động gửi đơn từ bàn ăn sẽ ngay lập tức kích hoạt chuông báo và hiển thị trực quan trên màn hình của bếp. Khi đầu bếp hoàn thành chế biến và nhấn 'Xong món', màn hình của phục vụ phụ trách khu vực bàn đó sẽ ngay lập tức nhấp nháy báo hiệu món ăn sẵn sàng bưng ra.")
    
    add_heading_styled(doc, "3.3. Trợ lý AI bán chéo sản phẩm (Gemini AI integration)", level=2)
    add_paragraph_styled(doc, "Tích hợp Gemini AI giúp phân tích thông tin các món ăn hiện tại khách hàng đang dùng tại bàn. Trợ lý AI sẽ gợi ý các món ăn kèm hoặc đồ uống phù hợp nhất (ví dụ: khách ăn lẩu thái cay thì gợi ý nước dừa xiêm hoặc trà đá mát lạnh) để nhân viên phục vụ khéo léo mời chào khách hàng, nâng cao doanh thu.")
    
    add_heading_styled(doc, "3.4. Quét mã VietQR động để thanh toán nhanh", level=2)
    add_paragraph_styled(doc, "Hệ thống tự động liên kết với API VietQR để tạo mã QR thanh toán động cho từng bàn ăn. Mã QR này chứa sẵn: Số tài khoản chủ nhà hàng, số tiền chính xác cần thanh toán của hóa đơn và nội dung chuyển khoản chi tiết (Ví dụ: 'Thanh toan ban Ban 04'). Khách hàng chỉ cần quét mã bằng ứng dụng ngân hàng, tiền chuyển thẳng vào tài khoản của nhà hàng mà không lo nhầm lẫn số tiền hay số tài khoản.")

    add_heading_styled(doc, "4. Thiết kế các giao diện màn hình và biểu mẫu (Screens & Forms)", level=1)
    add_paragraph_styled(doc, "Hệ thống được phát triển theo mô hình Single Page Application (SPA) gồm các màn hình chính sau:")
    
    add_bullet_styled(doc, " Giao diện hiện đại hiển thị banner chương trình khuyến mãi, danh mục các món ăn nổi bật, các bài viết tin tức ẩm thực và mục gửi phản hồi/đánh giá từ thực khách.", bold_prefix="Màn hình Trang chủ Khách hàng (Customer Home):")
    add_bullet_styled(doc, " Hiển thị danh mục món ăn phân loại rõ ràng (Món chính, Món khai vị, Đồ uống), hỗ trợ tìm kiếm món ăn nhanh và giỏ hàng nổi để khách chọn món, ghi chú yêu cầu đặc biệt gửi xuống bếp.", bold_prefix="Màn hình Menu Gọi Món & Giỏ Hàng:")
    add_bullet_styled(doc, " Biểu mẫu điền thông tin họ tên, số điện thoại, chọn ngày giờ hẹn trước và chọn vị trí bàn mong muốn.", bold_prefix="Biểu mẫu Đặt bàn hẹn giờ (Booking Form):")
    add_bullet_styled(doc, " Hiển thị sơ đồ bàn ăn được phân màu theo trạng thái (Xanh: Trống, Vàng: Khách đặt cọc trước, Đỏ: Đang có khách ăn, Tím: Khách vừa đi và cần dọn dẹp). Tích hợp khu vực hiển thị danh sách các món bếp vừa nấu xong cần bưng ra bàn khẩn cấp.", bold_prefix="Màn hình Vận hành của Phục Vụ (Waiter Dashboard):")
    add_bullet_styled(doc, " Giao diện tối giản tối ưu hiển thị danh sách các món ăn cần chế biến, sắp xếp theo thời gian gửi đơn từ cũ đến mới, hiển thị đầy đủ ghi chú món ăn của khách.", bold_prefix="Màn hình Chế biến của Bếp (Kitchen Dashboard):")
    add_bullet_styled(doc, " Bảng điều khiển trung tâm hiển thị các con số thống kê doanh thu, số đơn thành công. Tích hợp các biểu đồ trực quan (Biểu đồ doanh thu 7 ngày qua và biểu đồ Top 5 món bán chạy nhất) nhờ Chart.js.", bold_prefix="Màn hình Admin Dashboard (Analytics):")
    add_bullet_styled(doc, " Nơi Admin quản lý tài khoản nhân viên, chỉnh sửa thực đơn món ăn, quản lý nhập kho nguyên liệu và thiết lập định lượng công thức nấu ăn.", bold_prefix="Các Màn hình Quản lý Danh mục (CRUD Screens):")
    
    save_document(doc, "1_Mo_Ta_He_Thong.docx")

# ==============================================================================
# FILE 2: USE CASE (SoDoUsecase.docx)
# ==============================================================================
def generate_file_usecase():
    doc = setup_document_base("PHÂN TÍCH SƠ ĐỒ USE CASE HỆ THỐNG\nHỆ THỐNG QUẢN LÝ NHÀ HÀNG MỘC VỊ RESTAURANT")
    
    add_heading_styled(doc, "1. Tổng quan về các tác nhân (Actors) trong hệ thống", level=1)
    add_paragraph_styled(doc, "Sơ đồ Use Case là công cụ trực quan hóa mối quan hệ tương tác giữa người dùng và các chức năng của hệ thống. Hệ thống Mộc Vị Restaurant có 4 tác nhân chính tương tác với các Use Case chuyên biệt:")
    
    add_bullet_styled(doc, " Là người trực tiếp trải nghiệm ẩm thực, tương tác qua giao diện web công cộng.", bold_prefix="Khách hàng (Customer):")
    add_bullet_styled(doc, " Nhân viên tiền sảnh chịu trách nhiệm hỗ trợ khách, điều phối bàn ăn và thanh toán hóa đơn.", bold_prefix="Phục vụ (Waiter):")
    add_bullet_styled(doc, " Nhân viên chế biến chịu trách nhiệm tiếp nhận món ăn và quản lý trạng thái món.", bold_prefix="Bếp (Kitchen):")
    add_bullet_styled(doc, " Quản lý cấp cao phụ trách cấu hình hệ thống, quản lý tài sản, nhân sự và theo dõi báo cáo kinh doanh.", bold_prefix="Quản trị / Admin (Admin/Manager):")
    
    add_heading_styled(doc, "2. Sơ đồ Use Case tổng quát hệ thống", level=1)
    add_paragraph_styled(doc, "Hệ thống được thiết kế theo cấu trúc module hóa phân quyền rõ rệt. Dưới đây là phân nhóm các Use Case theo từng tác nhân:")
    
    add_heading_styled(doc, "2.1. Phân hệ Khách hàng (Customer Use Cases)", level=2)
    add_bullet_styled(doc, "Xem danh sách món ăn & tìm kiếm món ăn")
    add_bullet_styled(doc, "Đặt bàn hẹn giờ (cung cấp thông tin ngày giờ, số khách)")
    add_bullet_styled(doc, "Gọi món tại bàn (Order) qua mã QR")
    add_bullet_styled(doc, "Xem chi tiết giỏ hàng và theo dõi đơn hàng đang nấu")
    add_bullet_styled(doc, "Gửi đánh giá món ăn (Review)")
    add_bullet_styled(doc, "Đăng ký thành viên & tích điểm loyalty tích lũy")
    
    add_heading_styled(doc, "2.2. Phân hệ Nhân viên Phục vụ (Waiter Use Cases)", level=2)
    add_bullet_styled(doc, "Xem sơ đồ trạng thái bàn ăn (real-time)")
    add_bullet_styled(doc, "Hỗ trợ khách gọi thêm món tại bàn")
    add_bullet_styled(doc, "Chuyển bàn ăn tự động (Chuyển đơn hàng sang bàn trống mới)")
    add_bullet_styled(doc, "Xác nhận đã phục vụ / bưng món ra bàn (Update Order Status = 3)")
    add_bullet_styled(doc, "Tạo hóa đơn tạm tính và in hóa đơn")
    add_bullet_styled(doc, "Tham khảo gợi ý bán chéo món ăn từ AI chatbot")
    
    add_heading_styled(doc, "2.3. Phân hệ Nhân viên Bếp (Kitchen Use Cases)", level=2)
    add_bullet_styled(doc, "Nhận thông báo chế biến món ăn mới (real-time)")
    add_bullet_styled(doc, "Cập nhật trạng thái chế biến món (Đang nấu -> Đã xong)")
    add_bullet_styled(doc, "Báo hết món ăn trực tiếp trên thực đơn (Ngưng nhận đơn món này)")
    add_bullet_styled(doc, "Xem mức tồn kho nguyên liệu phục vụ nấu nướng")
    
    add_heading_styled(doc, "2.4. Phân hệ Quản trị viên (Admin Use Cases)", level=2)
    add_bullet_styled(doc, "Quản lý nhân viên (Thêm, sửa, xóa tài khoản và cấu hình vai trò)")
    add_bullet_styled(doc, "Quản lý thực đơn (CRUD món ăn, danh mục, cập nhật giá)")
    add_bullet_styled(doc, "Quản lý kho hàng & Công thức định lượng (Recipes)")
    add_bullet_styled(doc, "Xem báo cáo thống kê doanh thu và phân tích hiệu quả kinh doanh")
    add_bullet_styled(doc, "Quản lý tin tức sự kiện và hồ sơ ứng tuyển tuyển dụng")

    add_heading_styled(doc, "3. Đặc tả chi tiết các Use Case quan trọng", level=1)
    add_paragraph_styled(doc, "Dưới đây là đặc tả chi tiết của một số Use Case then chốt làm nên thế mạnh đột phá của dự án:")
    
    # Tạo bảng Use case đặc tả
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Tên Use Case'
    hdr_cells[1].text = 'Tác nhân chính'
    hdr_cells[2].text = 'Mô tả tóm tắt luồng hoạt động chính'
    
    # Set background cho header table
    for cell in hdr_cells:
        set_cell_background(cell, "1E3A8A")
        set_cell_margins(cell)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

    uc_data = [
        ("Gọi món tại bàn qua QR", "Khách hàng", "Khách quét mã QR tại bàn -> Trình duyệt hiển thị thực đơn -> Khách chọn món và gửi -> Hệ thống lưu đơn trạng thái 'Chờ chế biến' -> Gửi tín hiệu real-time xuống Bếp."),
        ("Đồng bộ chế biến tại Bếp", "Nhân viên Bếp", "Bếp nhận thông tin món ăn trên màn hình -> Click chọn 'Bắt đầu nấu' -> Click chọn 'Xong món' -> Hệ thống tự động trừ nguyên liệu tồn kho tương ứng và gửi thông báo chuông tới Phục vụ."),
        ("Bán chéo món ăn bằng AI", "Nhân viên Phục vụ", "Phục vụ mở modal chi tiết bàn ăn -> Click 'AI gợi ý mời món' -> Hệ thống gửi danh sách món bàn đó đang ăn lên Gemini AI -> AI trả về kịch bản gợi ý món phụ hợp -> Phục vụ ra mời khách."),
        ("Thanh toán động VietQR", "Phục vụ & Khách", "Phục vụ click in hóa đơn tạm tính -> Hệ thống tạo VietQR động chứa Số tiền hóa đơn và số bàn -> Khách quét mã để thanh toán -> Hệ thống cập nhật bàn ăn thành 'Cần dọn dẹp'.")
    ]
    
    for name, actor, desc in uc_data:
        row_cells = table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = actor
        row_cells[2].text = desc
        for cell in row_cells:
            set_cell_margins(cell)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
                    
    save_document(doc, "2_Use_Case.docx")

# ==============================================================================
# FILE 3: THIẾT KẾ CƠ SỞ DỮ LIỆU (ThietKeDatabase.docx)
# ==============================================================================
def generate_file_database():
    doc = setup_document_base("TÀI LIỆU THIẾT KẾ CƠ SỞ DỮ LIỆU CHI TIẾT\nHỆ THỐNG QUẢN LÝ NHÀ HÀNG MỘC VỊ RESTAURANT")
    
    add_heading_styled(doc, "1. Tổng quan kiến trúc Database", level=1)
    add_paragraph_styled(doc, "Cơ sở dữ liệu của dự án Mộc Vị Restaurant được xây dựng trên hệ quản trị cơ sở dữ liệu Microsoft SQL Server. Cấu trúc bảng được chuẩn hóa tối ưu đến dạng chuẩn 3 (3NF) nhằm hạn chế tối đa việc dư thừa dữ liệu, đồng thời thiết lập đầy đủ các ràng buộc toàn vẹn dữ liệu (Khóa chính, Khóa ngoại, Check Constraint) để bảo vệ tính chính xác của dữ liệu nghiệp vụ nhà hàng.")
    add_paragraph_styled(doc, "Hệ thống bao gồm tổng cộng 13 bảng liên kết chặt chẽ với nhau, bao quát toàn bộ hoạt động từ tài khoản, thực đơn, giao dịch hóa đơn, quản lý kho cho tới các tương tác cộng thêm như bài đăng tuyển dụng, khuyến mãi voucher.")
    
    add_heading_styled(doc, "2. Danh sách và đặc tả các bảng chi tiết trong hệ thống", level=1)
    add_paragraph_styled(doc, "Dưới đây là đặc tả chi tiết cấu trúc cột, kiểu dữ liệu và ràng buộc khóa của từng bảng dữ liệu:")
    
    # Hàm vẽ bảng đặc tả table
    def create_table_spec(doc, table_name, description, columns):
        add_heading_styled(doc, f"Bảng: {table_name} ({description})", level=2)
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Tên Cột'
        hdr_cells[1].text = 'Kiểu dữ liệu'
        hdr_cells[2].text = 'Khóa'
        hdr_cells[3].text = 'Cho phép Null'
        hdr_cells[4].text = 'Mô tả ý nghĩa nghiệp vụ'
        
        for cell in hdr_cells:
            set_cell_background(cell, "334155") # Slate background
            set_cell_margins(cell)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
                    
        for col_name, col_type, key, is_null, desc in columns:
            row_cells = table.add_row().cells
            row_cells[0].text = col_name
            row_cells[1].text = col_type
            row_cells[2].text = key
            row_cells[3].text = is_null
            row_cells[4].text = desc
            for cell in row_cells:
                set_cell_margins(cell)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)
        doc.add_paragraph() # Dòng trống ngăn cách
        
    # Đặc tả 1: Accounts
    create_table_spec(doc, "Accounts", "Quản lý thông tin tài khoản người dùng và khách hàng", [
        ("username", "VARCHAR(50)", "PK", "No", "Tên đăng nhập duy nhất của tài khoản"),
        ("password", "VARCHAR(100)", "-", "No", "Mật khẩu mã hóa"),
        ("fullname", "NVARCHAR(200)", "-", "No", "Họ và tên đầy đủ"),
        ("email", "VARCHAR(100)", "-", "No", "Địa chỉ email liên lạc"),
        ("photo", "VARCHAR(255)", "-", "Yes", "Đường dẫn ảnh đại diện"),
        ("total_spent", "FLOAT", "-", "Yes", "Tổng số tiền đã tích lũy tiêu dùng (để xét hạng thành viên)"),
        ("loyalty_points", "INT", "-", "Yes", "Số điểm tích lũy thành viên hiện tại"),
        ("tier", "VARCHAR(20)", "-", "Yes", "Hạng thành viên (BRONZE, SILVER, GOLD, DIAMOND)")
    ])
    
    # Đặc tả 2: Products
    create_table_spec(doc, "Products", "Quản lý danh sách các món ăn trong thực đơn", [
        ("id", "INT (IDENTITY)", "PK", "No", "Mã món ăn tự tăng"),
        ("name", "NVARCHAR(200)", "-", "No", "Tên món ăn (Ví dụ: Lẩu Thái Hải Sản)"),
        ("price", "FLOAT", "-", "No", "Giá bán của món ăn"),
        ("image", "VARCHAR(255)", "-", "Yes", "Đường dẫn hình ảnh minh họa món ăn"),
        ("description", "NVARCHAR(MAX)", "-", "Yes", "Mô tả chi tiết món ăn (nguyên liệu nổi bật...)"),
        ("create_date", "DATE", "-", "Yes", "Ngày cập nhật món ăn lên hệ thống"),
        ("available", "BIT", "-", "Yes", "Trạng thái còn bán món này hay không (1: Bán, 0: Ngừng)"),
        ("category_id", "INT", "FK", "Yes", "Liên kết với danh mục sản phẩm (Categories)")
    ])

    # Đặc tả 3: Orders
    create_table_spec(doc, "Orders", "Quản lý thông tin chung của đơn đặt bàn / đặt món", [
        ("id", "INT (IDENTITY)", "PK", "No", "Mã đơn hàng duy nhất"),
        ("create_date", "DATETIME", "-", "Yes", "Thời gian khởi tạo đơn hàng"),
        ("address", "NVARCHAR(500)", "-", "Yes", "Chứa thông tin số bàn ăn (Ví dụ: 'Bàn: Bàn 03') hoặc địa chỉ giao hàng"),
        ("status", "INT", "-", "Yes", "Trạng thái đơn: 1: Đang nấu, 2: Xong món, 3: Đang ăn, 4: Đã thanh toán, 5: Đơn hẹn trước"),
        ("note", "NVARCHAR(MAX)", "-", "Yes", "Ghi chú món ăn đặc biệt của khách (không cay, ít hành...)"),
        ("voucher_code", "VARCHAR(50)", "-", "Yes", "Mã voucher áp dụng giảm giá nếu có"),
        ("username", "VARCHAR(50)", "FK", "Yes", "Mã tài khoản khách hàng đặt đơn")
    ])

    # Đặc tả 4: OrderDetails
    create_table_spec(doc, "OrderDetails", "Chi tiết các món ăn và số lượng trong từng hóa đơn đơn hàng", [
        ("id", "INT (IDENTITY)", "PK", "No", "Mã chi tiết đơn hàng"),
        ("price", "FLOAT", "-", "Yes", "Đơn giá món ăn tại thời điểm bán"),
        ("quantity", "INT", "-", "Yes", "Số lượng món ăn đặt mua"),
        ("product_id", "INT", "FK", "Yes", "Mã liên kết tới sản phẩm"),
        ("order_id", "INT", "FK", "Yes", "Mã liên kết tới đơn hàng tổng")
    ])

    # Đặc tả 5: Ingredients & Recipes
    create_table_spec(doc, "Ingredients", "Quản lý nguyên liệu thô tồn kho trong bếp", [
        ("id", "BIGINT (IDENTITY)", "PK", "No", "Mã nguyên liệu thô"),
        ("name", "NVARCHAR(200)", "-", "No", "Tên nguyên liệu (Ví dụ: Thịt Bò Mỹ, Tôm Sú)"),
        ("quantity", "FLOAT", "-", "Yes", "Trọng lượng / số lượng còn lại trong kho bếp"),
        ("unit", "NVARCHAR(50)", "-", "Yes", "Đơn vị tính nguyên liệu (kg, gram, lít, hộp)"),
        ("min_stock", "FLOAT", "-", "Yes", "Mức tồn tối thiểu để phát cảnh báo hết hàng")
    ])
    
    create_table_spec(doc, "Recipes", "Công thức định lượng hao hụt nguyên liệu thô khi chế biến một món ăn", [
        ("id", "BIGINT (IDENTITY)", "PK", "No", "Mã định lượng công thức"),
        ("product_id", "INT", "FK", "Yes", "Mã món ăn thành phẩm"),
        ("ingredient_id", "BIGINT", "FK", "Yes", "Mã nguyên liệu cần dùng hao phí"),
        ("amount_required", "FLOAT", "-", "Yes", "Khối lượng nguyên liệu tiêu thụ cho 1 phần ăn")
    ])

    # Đặc tả 6: restaurant_table
    create_table_spec(doc, "restaurant_table", "Quản lý sơ đồ và trạng thái các bàn ăn ở tiền sảnh", [
        ("id", "INT (IDENTITY)", "PK", "No", "Mã định danh bàn ăn"),
        ("name", "NVARCHAR(255)", "-", "Yes", "Tên số bàn (Ví dụ: Bàn 01, Bàn 02)"),
        ("floor", "NVARCHAR(255)", "-", "Yes", "Tầng vị trí bàn (Tầng 1, Tầng 2, Sân vườn)"),
        ("is_occupied", "INT", "-", "Yes", "Trạng thái thực tế: 0: Trống, 1: Đã đặt cọc, 2: Có khách ăn, 3: Chờ dọn bàn"),
        ("has_view", "BIT", "-", "Yes", "Bàn ăn có view đẹp ngoài trời hay không"),
        ("reserved_time", "VARCHAR(255)", "-", "Yes", "Mốc thời gian khách hẹn trước"),
        ("capacity", "INT", "-", "Yes", "Sức chứa tối đa số ghế tại bàn"),
        ("view_type", "NVARCHAR(50)", "-", "Yes", "Loại view cảnh quan (Cửa sổ, Ban công, Bể cá...)")
    ])
    
    save_document(doc, "3_Database_Design.docx")

# ==============================================================================
# FILE 4: CÔNG NGHỆ DỰ ÁN (CongNgheDuAn.docx)
# ==============================================================================
def generate_file_congnghe():
    doc = setup_document_base("BÁO CÁO CÔNG NGHỆ VÀ THIẾT KẾ KIẾN TRÚC HỆ THỐNG\nHỆ THỐNG QUẢN LÝ NHÀ HÀNG MỘC VỊ RESTAURANT")
    
    add_heading_styled(doc, "1. Mô hình kiến trúc tổng quan hệ thống", level=1)
    add_paragraph_styled(doc, "Dự án Mộc Vị Restaurant được xây dựng trên kiến trúc Single Page Application (SPA) hiện đại kết hợp mô hình Service-Oriented Architecture (SOA). Sự tách biệt hoàn toàn giữa giao diện người dùng (Frontend) và logic xử lý nghiệp vụ (Backend) giúp hệ thống đạt hiệu năng cao, tối ưu băng thông mạng và mang lại trải nghiệm mượt mà không độ trễ cho người dùng.")
    
    add_paragraph_styled(doc, "Luồng trao đổi dữ liệu chính trong hệ thống:")
    add_bullet_styled(doc, " Giao diện Frontend tương tác bất đồng bộ thông qua các hàm gọi API RESTful (định dạng dữ liệu JSON). Điều này giúp trình duyệt không cần tải lại toàn bộ trang web khi thực hiện các tác vụ cập nhật dữ liệu.", bold_prefix="RESTful API (JSON):")
    add_bullet_styled(doc, " Đối với các tác vụ yêu cầu đồng bộ tức thời như thông báo có món ăn mới từ Bếp đến Phục vụ, hệ thống thiết lập kết nối song công liên tục sử dụng giao thức WebSocket STOMP bảo đảm dữ liệu truyền đi trong vài mili-giây.", bold_prefix="Real-time WebSockets (TCP):")
    add_bullet_styled(doc, " Toàn bộ các API được che chắn và kiểm tra phân quyền chặt chẽ theo từng role nhờ bộ lọc Security Filter Chain và mã hóa Token bảo mật JWT.", bold_prefix="State-less Security:")
    
    add_heading_styled(doc, "2. Công nghệ phát triển phía Backend (Spring Boot)", level=1)
    add_paragraph_styled(doc, "Backend đóng vai trò 'đầu não' điều hành toàn bộ logic tính toán và quản trị dữ liệu. Các công nghệ cốt lõi được sử dụng bao gồm:")
    
    add_bullet_styled(doc, " Cung cấp bộ khung cấu hình nhanh (Convention over Configuration), giúp khởi chạy ứng dụng nhanh với Web Server nhúng Tomcat.", bold_prefix="Spring Boot (Java 21):")
    add_bullet_styled(doc, " Bộ lọc phân quyền chặt chẽ ở cấp độ Endpoint và Method (PreAuthorize), mã hóa mật khẩu người dùng an toàn.", bold_prefix="Spring Security:")
    add_bullet_styled(doc, " Cơ chế bảo mật phi trạng thái (Stateless), tạo chuỗi ký số JWT gửi kèm header mỗi yêu cầu của client, hạn chế tối đa việc chiếm dụng Session trên server.", bold_prefix="JSON Web Token (JWT):")
    add_bullet_styled(doc, " Ánh xạ thực thể cơ sở dữ liệu thành đối tượng Java (Object-Relational Mapping), hỗ trợ truy vấn CSDL cực kỳ an toàn, chống lỗ hổng SQL Injection phổ biến.", bold_prefix="Spring Data JPA & Hibernate:")
    add_bullet_styled(doc, " Thiết lập Broker trung gian kết nối các tin nhắn và phân phát tới các kênh đăng ký (/topic/kitchen, /topic/waiter).", bold_prefix="Spring WebSocket (STOMP):")
    add_bullet_styled(doc, " Tích hợp SDK AI thế hệ mới nhất của Google để gửi yêu cầu phân tích hóa đơn và đưa ra đề xuất bán thêm sản phẩm hiệu quả theo thời gian thực.", bold_prefix="Gemini AI Integration:")
    
    add_heading_styled(doc, "3. Công nghệ phát triển phía Frontend (Vue 3 / Vite)", level=1)
    add_paragraph_styled(doc, "Giao diện Frontend được thiết kế tối giản, tập trung vào trải nghiệm trực quan sống động của nhân viên và khách hàng:")
    
    add_bullet_styled(doc, " Sử dụng cú pháp Composition API (<script setup>) hiện đại, mang lại khả năng tái sử dụng mã nguồn xuất sắc, khả năng binding dữ liệu hai chiều (two-way data binding) phản hồi tức thì cực kỳ mượt mà.", bold_prefix="Vue 3 (Composition API):")
    add_bullet_styled(doc, " Công cụ build ứng dụng thế hệ mới, thay thế Webpack cũ kỹ, mang lại tốc độ chạy thử máy chủ cục bộ (Hot Module Replacement) chỉ trong chưa đầy 1 giây.", bold_prefix="Vite Build Tool:")
    add_bullet_styled(doc, " Thư viện quản lý State tập trung, dùng để đồng bộ giỏ hàng, thông tin đăng nhập và token JWT của người dùng xuyên suốt toàn bộ ứng dụng.", bold_prefix="Pinia State Management:")
    add_bullet_styled(doc, " Thư viện kết nối HTTP mạnh mẽ, hỗ trợ Interceptor để tự động đính kèm Token JWT vào header mỗi khi gọi API lên Backend.", bold_prefix="Axios HTTP Client:")
    add_bullet_styled(doc, " Kết nối và duy trì kết nối WebSocket liên tục, tự động đăng ký vào các topic thông báo của máy chủ.", bold_prefix="SockJS & StompJS:")
    add_bullet_styled(doc, " Vẽ biểu đồ phân tích doanh thu trực quan ngay trên trình duyệt.", bold_prefix="Chart.js & Vue-chartjs:")
    add_bullet_styled(doc, " Hỗ trợ đa ngôn ngữ linh hoạt cho thực khách quốc tế.", bold_prefix="Vue-i18n:")

    add_heading_styled(doc, "4. Cơ sở dữ liệu và hạ tầng lưu trữ", level=1)
    add_bullet_styled(doc, " Hệ quản trị cơ sở dữ liệu quan hệ mạnh mẽ, đảm bảo tính toàn vẹn cao, hỗ trợ đắc lực cho các truy vấn phức tạp liên quan đến tính toán doanh thu và quản lý định lượng kho hàng.", bold_prefix="Hệ quản trị CSDL: Microsoft SQL Server:")
    add_bullet_styled(doc, " Máy chủ ảo Windows OS chạy trực tiếp máy chủ Spring Boot, máy chủ chứa web tĩnh Frontend và máy chủ SQL Server liên kết an toàn cục bộ.", bold_prefix="Hạ tầng chạy thử:")
    
    save_document(doc, "4_Cong_Nghe_Du_An.docx")

# ==============================================================================
# CHƯƠNG TRÌNH CHÍNH
# ==============================================================================
if __name__ == "__main__":
    print("--- BẮT ĐẦU TẠO CÁC FILE WORD HƯỚNG DẪN ĐỒ ÁN TỐT NGHIỆP ---")
    generate_file_mota()
    generate_file_usecase()
    generate_file_database()
    generate_file_congnghe()
    print("--- HOÀN THÀNH TẤT CẢ 4 FILE DOCX CHO ĐỒ ÁN! ---")
